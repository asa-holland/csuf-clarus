import chardet
import pdfminer.high_level
import pdfminer.layout
from docx import Document
from bs4 import BeautifulSoup
import readability

ALLOWED_EXTENSIONS = {"pdf", "docx", "html", "htm", "txt"}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_pdf_text(file_path):
    try:
        laparams = pdfminer.layout.LAParams()
        laparams.all_texts = True

        with open(file_path, "rb") as file:
            text = pdfminer.high_level.extract_text(file, laparams=laparams)

        metadata = {
            "format": "PDF",
            "extraction_method": "pdfminer.six",
            "encoding": "UTF-8",
        }

        return text, metadata
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")


def extract_docx_text(file_path):
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        core_props = doc.core_properties
        metadata = {
            "format": "DOCX",
            "extraction_method": "python-docx",
            "encoding": "UTF-8",
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }

        return text, metadata
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")


def extract_html_text(file_path):
    try:
        with open(file_path, "rb") as file:
            raw_html = file.read()

        detected = chardet.detect(raw_html)
        encoding = detected.get("encoding", "utf-8")

        try:
            html_content = raw_html.decode(encoding)
        except UnicodeDecodeError:
            html_content = raw_html.decode("utf-8", errors="ignore")

        readable_document = readability.Document(html_content)
        clean_html = readable_document.summary()

        soup = BeautifulSoup(clean_html, "html.parser")
        text = soup.get_text(separator="\n", strip=True)

        title = soup.find("title")
        title_text = title.get_text().strip() if title else ""

        metadata = {
            "format": "HTML",
            "extraction_method": "readability-lxml + BeautifulSoup",
            "encoding": encoding,
            "title": title_text,
            "readability_score": readable_document.score(),
        }

        return text, metadata
    except Exception as e:
        raise Exception(f"HTML extraction failed: {str(e)}")


def extract_txt_text(file_path):
    try:
        with open(file_path, "rb") as file:
            raw_content = file.read()

        detected = chardet.detect(raw_content)
        encoding = detected.get("encoding", "utf-8")
        confidence = detected.get("confidence", 0)

        try:
            text = raw_content.decode(encoding)
        except UnicodeDecodeError:
            text = raw_content.decode("utf-8", errors="ignore")
            encoding = "utf-8 (fallback)"

        lines = text.split("\n")
        words = text.split()

        metadata = {
            "format": "TXT",
            "extraction_method": "direct read + chardet",
            "encoding": encoding,
            "encoding_confidence": confidence,
            "line_count": len(lines),
            "word_count": len(words),
            "character_count": len(text),
        }

        return text, metadata
    except Exception as e:
        raise Exception(f"TXT extraction failed: {str(e)}")


def extract_text_from_file(file_path, file_extension):
    extractors = {
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "html": extract_html_text,
        "htm": extract_html_text,
        "txt": extract_txt_text,
    }

    extractor = extractors.get(file_extension.lower())
    if not extractor:
        raise Exception(f"Unsupported file format: {file_extension}")

    return extractor(file_path)
